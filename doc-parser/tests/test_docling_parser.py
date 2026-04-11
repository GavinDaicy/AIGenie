"""
I2-B7 DoclingParser 测试。

分两类：
1. 使用 Mock 的单元测试 — 无需安装 Docling，验证解析逻辑
2. 集成测试（skipif Docling 不可用）— 真实解析 DOCX / PDF fixture 文件
"""
import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock

import pytest

# ---------- Docling 可用性检查 ----------
try:
    import docling  # noqa: F401
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False

FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ============================================================
# 辅助：构造 Docling Mock 文档
# ============================================================

def _make_prov(page_no=1, l=72.0, t=700.0, r=540.0, b=680.0):
    """构造 ProvenanceItem mock，bbox 使用 BOTTOMLEFT 坐标。"""
    bbox = MagicMock()
    bbox.l = l
    bbox.t = t
    bbox.r = r
    bbox.b = b
    bbox.coord_origin = MagicMock()
    bbox.coord_origin.__str__ = lambda self: "CoordOrigin.BOTTOMLEFT"

    # to_top_left_origin 模拟：new_t = page_h - t, new_b = page_h - b
    def to_top_left(page_h):
        converted = MagicMock()
        converted.l = l
        converted.t = page_h - t   # 顶边（小 y）
        converted.r = r
        converted.b = page_h - b   # 底边（大 y）
        return converted

    bbox.to_top_left_origin = to_top_left

    prov = MagicMock()
    prov.page_no = page_no
    prov.bbox = bbox
    return prov


def _make_item(label_val, text, prov=None, level=None):
    item = MagicMock()
    label = MagicMock()
    label.value = label_val
    item.label = label
    item.text = text
    item.prov = [prov] if prov else []
    if level is not None:
        item.level = level
    return item


def _make_table_item(text, prov=None):
    item = _make_item("table", text, prov)
    item.export_to_markdown.return_value = text
    return item


def _make_mock_doc(items, page_sizes=None):
    """构造 DoclingDocument mock。"""
    doc = MagicMock()

    # pages
    pages = {}
    for page_no, (w, h) in (page_sizes or {1: (595.3, 841.9)}).items():
        page = MagicMock()
        page.size = MagicMock()
        page.size.width = w
        page.size.height = h
        pages[page_no] = page
    doc.pages = pages

    doc.iterate_items.return_value = items
    return doc


# ============================================================
# 单元测试（Mock Docling）
# ============================================================

class TestDoclingParserUnit:

    def _parse_with_mock(self, items, page_sizes=None):
        from app.parsers.docling_parser import DoclingParser

        mock_doc = _make_mock_doc(items, page_sizes)
        mock_result = MagicMock()
        mock_result.document = mock_doc

        mock_converter = MagicMock()
        mock_converter.convert.return_value = mock_result

        parser = DoclingParser()
        parser._get_converter = lambda: mock_converter
        return parser.parse("/fake/path/doc.docx", "job-test")

    def test_heading_block_created(self):
        prov = _make_prov(page_no=1)
        items = [(_make_item("title", "My Title", prov), 0)]
        output = self._parse_with_mock(items)

        assert len(output.blocks) == 1
        assert output.blocks[0].type.value == "heading"
        assert output.blocks[0].text == "My Title"
        assert output.blocks[0].level == 1

    def test_section_header_block(self):
        prov = _make_prov(page_no=1)
        items = [(_make_item("section_header", "Introduction", prov, level=1), 1)]
        output = self._parse_with_mock(items)

        assert output.blocks[0].type.value == "heading"
        assert output.blocks[0].text == "Introduction"

    def test_paragraph_block_created(self):
        prov = _make_prov(page_no=1)
        items = [(_make_item("text", "Some paragraph text.", prov), 1)]
        output = self._parse_with_mock(items)

        assert len(output.blocks) == 1
        assert output.blocks[0].type.value == "paragraph"

    def test_table_block_created(self):
        prov = _make_prov(page_no=1)
        md_table = "| A | B |\n|---|---|\n| 1 | 2 |"
        items = [(_make_table_item(md_table, prov), 1)]
        output = self._parse_with_mock(items)

        assert output.blocks[0].type.value == "table"
        assert "| A | B |" in output.blocks[0].text

    def test_list_item_block(self):
        prov = _make_prov(page_no=1)
        items = [(_make_item("list_item", "Item one", prov), 1)]
        output = self._parse_with_mock(items)

        assert output.blocks[0].type.value == "list"
        assert "Item one" in output.blocks[0].text

    def test_image_block(self):
        prov = _make_prov(page_no=2)
        pic_item = _make_item("picture", "Figure caption", prov)
        pic_item.captions = []
        items = [(pic_item, 1)]
        output = self._parse_with_mock(items)

        assert output.blocks[0].type.value == "image"
        assert output.blocks[0].image_path is not None
        assert "img-" in output.blocks[0].image_path

    def test_page_header_is_skipped(self):
        prov = _make_prov()
        items = [(_make_item("page_header", "Header", prov), 0)]
        output = self._parse_with_mock(items)
        assert len(output.blocks) == 0

    def test_page_footer_is_skipped(self):
        prov = _make_prov()
        items = [(_make_item("page_footer", "Footer", prov), 0)]
        output = self._parse_with_mock(items)
        assert len(output.blocks) == 0

    def test_empty_text_block_is_skipped(self):
        prov = _make_prov()
        items = [(_make_item("text", "   ", prov), 1)]
        output = self._parse_with_mock(items)
        assert len(output.blocks) == 0

    def test_block_ids_are_sequential(self):
        prov = _make_prov()
        items = [
            (_make_item("title", "T1", prov), 0),
            (_make_item("text", "P1", prov), 1),
            (_make_item("text", "P2", prov), 1),
        ]
        output = self._parse_with_mock(items)
        ids = [b.block_id for b in output.blocks]
        assert ids == ["b-0001", "b-0002", "b-0003"]

    def test_coordinate_conversion_topleft(self):
        """验证坐标转换后 y0 < y1（top-left 坐标系）。"""
        page_h = 841.9
        prov = _make_prov(page_no=1, l=72.0, t=700.0, r=540.0, b=650.0)
        items = [(_make_item("text", "Some text.", prov), 1)]
        output = self._parse_with_mock(
            items, page_sizes={1: (595.3, page_h)}
        )

        bbox = output.blocks[0].bbox
        x0, y0, x1, y1 = bbox
        assert x0 == pytest.approx(72.0)
        assert x1 == pytest.approx(540.0)
        assert y0 < y1, f"Expected y0 < y1, got y0={y0}, y1={y1}"
        assert y0 >= 0
        assert y1 <= page_h

    def test_heading_path_tracks_parent_headings(self):
        """正文 block 的 heading_path 应包含前面的标题文字。"""
        prov = _make_prov()
        items = [
            (_make_item("title", "Chapter 1", prov), 0),
            (_make_item("text", "Body text.", prov), 1),
        ]
        output = self._parse_with_mock(items)

        body_block = output.blocks[1]
        assert body_block.heading_path is not None
        assert "Chapter 1" in body_block.heading_path

    def test_markdown_contains_block_anchors(self):
        prov = _make_prov()
        items = [
            (_make_item("title", "Title", prov), 0),
            (_make_item("text", "Content.", prov), 1),
        ]
        output = self._parse_with_mock(items)

        assert "<!-- @block:b-0001 -->" in output.markdown
        assert "<!-- @block:b-0002 -->" in output.markdown

    def test_positions_json_fields(self):
        prov = _make_prov()
        items = [(_make_item("text", "Hello.", prov), 1)]
        output = self._parse_with_mock(items)

        pos = output.positions
        assert pos.parser == "docling"
        assert pos.coordinate_system == "top-left"
        assert pos.total_pages >= 1
        assert isinstance(pos.blocks, list)

    def test_multi_page_block_page_numbers(self):
        prov1 = _make_prov(page_no=1)
        prov2 = _make_prov(page_no=3)
        items = [
            (_make_item("text", "Page 1 text.", prov1), 1),
            (_make_item("text", "Page 3 text.", prov2), 1),
        ]
        output = self._parse_with_mock(
            items, page_sizes={1: (595.3, 841.9), 3: (595.3, 841.9)}
        )

        assert output.blocks[0].page == 1
        assert output.blocks[1].page == 3

    def test_code_block(self):
        prov = _make_prov()
        items = [(_make_item("code", "x = 1", prov), 1)]
        output = self._parse_with_mock(items)

        assert output.blocks[0].type.value == "code"

    def test_formula_block(self):
        prov = _make_prov()
        items = [(_make_item("formula", r"E=mc^2", prov), 1)]
        output = self._parse_with_mock(items)

        assert output.blocks[0].type.value == "formula"


# ============================================================
# 集成测试（需要真实 Docling + fixture 文件）
# ============================================================

@pytest.mark.skipif(not HAS_DOCLING, reason="docling not installed")
class TestDoclingParserIntegration:
    """
    真实解析 fixture 文件。
    如果 fixture 文件不存在则自动跳过对应测试。
    """

    def _fixture(self, name: str) -> str:
        path = str(FIXTURES_DIR / name)
        if not os.path.exists(path):
            pytest.skip(f"Fixture not found: {path}")
        return path

    def _create_docx_fixture(self) -> str:
        """尝试用 python-docx 在 fixtures 目录创建 sample.docx。"""
        try:
            from docx import Document
            path = str(FIXTURES_DIR / "sample.docx")
            if not os.path.exists(path):
                doc = Document()
                doc.add_heading("Sample Document", level=1)
                doc.add_paragraph("This is the introduction paragraph.")
                doc.add_heading("Section 1", level=2)
                doc.add_paragraph("Section content goes here.")
                table = doc.add_table(rows=2, cols=2)
                table.cell(0, 0).text = "Header A"
                table.cell(0, 1).text = "Header B"
                table.cell(1, 0).text = "Value 1"
                table.cell(1, 1).text = "Value 2"
                FIXTURES_DIR.mkdir(exist_ok=True)
                doc.save(path)
            return path
        except ImportError:
            pytest.skip("python-docx not available")

    def test_docx_contains_block_anchors(self):
        path = self._create_docx_fixture()
        from app.parsers.docling_parser import DoclingParser
        output = DoclingParser().parse(path, "job-docx")
        assert "<!-- @block:b-0001 -->" in output.markdown

    def test_docx_positions_has_blocks(self):
        path = self._create_docx_fixture()
        from app.parsers.docling_parser import DoclingParser
        output = DoclingParser().parse(path, "job-docx2")
        assert len(output.positions.blocks) > 0

    def test_docx_blocks_have_text(self):
        path = self._create_docx_fixture()
        from app.parsers.docling_parser import DoclingParser
        output = DoclingParser().parse(path, "job-docx3")
        texts = [b.text for b in output.blocks]
        assert any("Sample Document" in t for t in texts)

    def _create_minimal_pdf(self) -> str:
        """
        用 reportlab（可选）或最小 PDF 字节串创建单页 PDF。
        如果两种方式都不可用则 skip。
        """
        path = str(FIXTURES_DIR / "minimal.pdf")
        if os.path.exists(path):
            return path

        # 尝试 reportlab
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            c = rl_canvas.Canvas(path)
            c.drawString(72, 700, "Hello World - Docling Test")
            c.drawString(72, 680, "This is a minimal test PDF.")
            c.save()
            return path
        except ImportError:
            pass

        # 回退：手写最小合法 PDF（单页，含一行文字）
        pdf_bytes = (
            b"%PDF-1.4\n"
            b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
            b" /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
            b"4 0 obj\n<< /Length 44 >>\nstream\n"
            b"BT /F1 12 Tf 72 700 Td (Hello Docling Test) Tj ET\n"
            b"endstream\nendobj\n"
            b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
            b"xref\n0 6\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000058 00000 n \n"
            b"0000000115 00000 n \n"
            b"0000000274 00000 n \n"
            b"0000000373 00000 n \n"
            b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
            b"startxref\n448\n%%EOF\n"
        )
        FIXTURES_DIR.mkdir(exist_ok=True)
        with open(path, "wb") as f:
            f.write(pdf_bytes)
        return path

    @pytest.mark.timeout(120)
    def test_pdf_has_page_and_bbox(self):
        path = self._create_minimal_pdf()
        from app.parsers.docling_parser import DoclingParser
        output = DoclingParser().parse(path, "job-pdf")
        for block in output.blocks:
            assert block.page >= 1
            assert len(block.bbox) == 4

    @pytest.mark.timeout(120)
    def test_pdf_coordinate_system_is_topleft(self):
        path = self._create_minimal_pdf()
        from app.parsers.docling_parser import DoclingParser
        output = DoclingParser().parse(path, "job-pdf2")
        assert output.positions.coordinate_system == "top-left"
        for block in output.blocks:
            x0, y0, x1, y1 = block.bbox
            if x0 != 0 or y0 != 0:
                assert y0 <= y1, f"Expected y0<=y1, got {y0}, {y1}"
